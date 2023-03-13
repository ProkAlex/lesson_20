'''
from docx import Document
from docx.shared import Inches

document = Document()
document.add_heading('Заголовок, ')
p = document.add_paragraph('Параграф, ')
p.add_run('Жирный, ').bold = True
p.add_run('Курсивный, ').italic = True
document.save('doca.doc')
file = open('text.txt')
text = file.read()
print(text, type(text), len(text))
'''
def lenfile(location):
    try:
        file = open(location)
        text = file.read()
        answer = len(text)
        return answer
    except FileNotFoundError:
        answer = 'file not found'
        return answer

print(lenfile('fil.txt'))

s = 'Мама мыла раму'

def firstword(stringa):
    word = ''
    for i in stringa:
        if i == ' ' or i == ',':
            break
        else:
            word += i
    return word
print(firstword(s))